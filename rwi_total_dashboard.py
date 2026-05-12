import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import os
import zipfile
import shutil
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches

# --- Page Configuration ---
st.set_page_config(
    page_title="RWI Full Baseline Report Dashboard",
    page_icon="📊",
    layout="wide"
)

# --- Custom Styling ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;700&display=swap');
    
    /* Premium Glassmorphism Theme */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    
    [data-testid="stSidebar"] {
        background-color: rgba(255, 255, 255, 0.5) !important;
        backdrop-filter: blur(10px);
        border-right: 1px solid rgba(255, 255, 255, 0.3);
    }
    
    .main-header {
        background: rgba(30, 61, 89, 0.85);
        backdrop-filter: blur(15px);
        color: white;
        padding: 50px;
        border-radius: 24px;
        text-align: center;
        margin-bottom: 40px;
        box-shadow: 0 20px 40px rgba(0,0,0,0.15);
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .section-card {
        background: rgba(255, 255, 255, 0.7);
        backdrop-filter: blur(12px);
        padding: 30px;
        border-radius: 20px;
        box-shadow: 0 8px 32px rgba(31, 38, 135, 0.07);
        margin-bottom: 40px;
        border: 1px solid rgba(255, 255, 255, 0.4);
        transition: transform 0.3s ease;
    }
    
    .section-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 40px rgba(31, 38, 135, 0.12);
    }
    
    .section-title {
        color: #1e3d59;
        font-size: 26px;
        font-weight: 700;
        margin-bottom: 25px;
        padding-bottom: 15px;
        border-bottom: 2px solid rgba(30, 61, 89, 0.1);
    }
    
    h1, h2, h3 { font-family: 'Outfit', sans-serif; font-weight: 700; }
    
    /* Table Styling */
    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
    }
    
    /* Button Styling */
    .stButton>button {
        background: linear-gradient(90deg, #1e3d59 0%, #2e7bcf 100%);
        color: white;
        border-radius: 12px;
        border: none;
        padding: 12px 24px;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        opacity: 0.9;
        transform: scale(1.02);
    }
    </style>
""", unsafe_allow_html=True)

# --- Constants ---
OUTPUT_DIR = "full_report_outputs"
COLORS = ["#636EFA", "#EF553B", "#00CC96", "#AB63FA", "#FFA15A", "#19D3F3"]

# --- Helper Functions ---
def setup_dirs():
    """Sets up the output directory with robust error handling for Windows file locks."""
    if os.path.exists(OUTPUT_DIR):
        try:
            shutil.rmtree(OUTPUT_DIR)
        except Exception:
            # If rmtree fails (common on Windows if files are open), try clearing files individually
            for filename in os.listdir(OUTPUT_DIR):
                file_path = os.path.join(OUTPUT_DIR, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception:
                    continue # Skip locked files
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def clean_df(df):
    """Cleans all object columns and headers in a dataframe."""
    df_clean = df.copy()
    for col in df_clean.columns:
        if df_clean[col].dtype == 'object':
            df_clean[col] = df_clean[col].apply(clean_underscore_values)
    df_clean.columns = [clean_underscore_values(c) for c in df_clean.columns]
    return df_clean

def save_and_render(section_id, title, df_table, fig):
    """
    Renders a standard section with a table and a chart in a side-by-side layout.
    Saves results to Excel and PNG for the final ZIP download.
    """
    # Remove section numbers from the chart title if present (e.g. "3.1. Households" -> "Households")
    clean_chart_title = title.split('. ', 2)[-1] if '. ' in title else title
    if clean_chart_title and clean_chart_title[0].isdigit() and '. ' in clean_chart_title:
        clean_chart_title = clean_chart_title.split('. ', 2)[-1]

    if fig:
        # Global Chart Styling
        fig.update_layout(
            title=clean_chart_title,
            colorway=COLORS,
            title_font=dict(size=20, family="Outfit", color="#1e3d59"),
            font=dict(family="Outfit"),
            legend=dict(
                orientation="v", 
                yanchor="middle", y=0.5, 
                xanchor="left", x=1.02
            ),
            xaxis=dict(categoryorder='total ascending'),
            margin=dict(r=120, t=80, b=40, l=40),
            plot_bgcolor="rgba(0,0,0,0)",
            paper_bgcolor="rgba(0,0,0,0)",
            height=450
        )
        
        # Enforce Percentage-Only Labels
        if isinstance(fig.data[0], (go.Pie,)):
            fig.update_traces(
                textinfo='percent', 
                textposition='inside',
                textfont=dict(size=14, color="white", weight='bold')
            )
        else:
            fig.update_traces(
                texttemplate='<b>%{y:.2f}%</b>', 
                textposition='outside',
                textfont=dict(weight='bold')
            )

    with st.container():
        st.markdown(f"<div class='section-card'><div class='section-title'>{section_id} - {title}</div>", unsafe_allow_html=True)
        c1, c2 = st.columns([1, 1.4])
        with c1:
            st.dataframe(df_table, use_container_width=True, hide_index=True)
            df_table.to_excel(os.path.join(OUTPUT_DIR, f"{section_id}.xlsx"), index=False)
        with c2:
            if fig:
                st.plotly_chart(fig, use_container_width=True, config={'displayModeBar': True})
                try: fig.write_image(os.path.join(OUTPUT_DIR, f"{section_id}.png"), scale=2)
                except: pass
        st.markdown("</div>", unsafe_allow_html=True) # Close section-card

def normalize_yes_no(series):
    """Standardizes Yes/No/None responses to lowercase 'yes' or 'no'."""
    return series.astype(str).str.lower().str.strip().replace(["nan", "none", "0", "0.0"], "no").replace(["1", "1.0"], "yes")

# ==========================================
# INDICATOR ANALYSIS (PORTED FROM SCRIPT)
# ==========================================

def run_3_households(df):
    # 3.1 Panchayath Level
    total_hh = len(df)
    gp_hh = pd.DataFrame([{
        "No of households": total_hh,
        "Households %": "100.00%"
    }])
    
    # 3.1 Chart - Simple single bar/pie representing the total
    fig_gp = px.pie(values=[total_hh], names=["Total Households"], title="3.1. Households – Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("3.1", "Households – Panchayath level", gp_hh, fig_gp)

    # 3.2 Village Level
    vil_hh = df.groupby("village")["farmer_name"].count().reset_index()
    vil_hh.columns = ["Village", "No of Households"]
    vil_hh["Households %"] = (vil_hh["No of Households"] / total_hh * 100).round(2)
    
    # Sort for chart
    vil_hh_sorted = vil_hh.sort_values("Households %", ascending=False)
    
    # Add Total Row
    total_row = pd.DataFrame([["Grand Total", total_hh, 100.0]], columns=vil_hh.columns)
    vil_hh_table = pd.concat([vil_hh, total_row], ignore_index=True)
    vil_hh_table["Households %"] = vil_hh_table["Households %"].astype(str) + "%"
    
    fig_vil = px.bar(vil_hh_sorted, x="Village", y="Households %", title="3.2. Households – Village level", color_discrete_sequence=COLORS)
    save_and_render("3.2", "Households – Village level", vil_hh_table, fig_vil)

def run_4_social_composition(df):
    # 4.1 Panchayath Level
    # No deduplication as per reference
    expected = ["SC", "ST", "OC", "BC"]
    caste_counts_ser = df["category"].value_counts()
    caste_counts = caste_counts_ser.reindex(expected, fill_value=0)
    # FORMULA ALIGNMENT: Denominator is sum of shown categories
    total_sum = caste_counts.sum()
    
    gp_caste_table = pd.DataFrame({
        "Category": expected + ["Total"],
        "No of households": caste_counts.tolist() + [total_sum],
        "Households %": [f"{(v/total_sum*100):.2f}%" if total_sum > 0 else "0.00%" for v in caste_counts] + ["100%"]
    })
    
    fig_gp = px.pie(names=expected, values=caste_counts.tolist(), title="4.1. Caste wise number of families - Panchayath level", hole=0.45, color_discrete_sequence=COLORS)
    save_and_render("4.1", "Caste wise number of families - Panchayath level", gp_caste_table, fig_gp)

    # 4.2 Village Level
    vil_caste = df.groupby(["village", "category"]).size().unstack(fill_value=0)
    vil_caste = vil_caste.reindex(columns=expected, fill_value=0)
    vil_caste["Total"] = vil_caste.sum(axis=1)
    
    table_df = vil_caste.reset_index()
    cols_order = ["village"]
    for c in expected:
        table_df[f"{c} %"] = (table_df[c] / table_df["Total"] * 100).round(2)
        cols_order.extend([c, f"{c} %"])
    cols_order.append("Total")
    
    # Grand Total Row
    gt_vals = table_df[expected + ["Total"]].sum()
    gt_row = {"village": "Grand Total"}
    for c in expected:
        gt_row[c] = gt_vals[c]
        gt_row[f"{c} %"] = round(gt_vals[c] / gt_vals["Total"] * 100, 2) if gt_vals["Total"] > 0 else 0.0
    gt_row["Total"] = gt_vals["Total"]
    table_df = pd.concat([table_df, pd.DataFrame([gt_row])], ignore_index=True)
    
    # Format percentages as strings for display
    for c in expected:
        table_df[f"{c} %"] = table_df[f"{c} %"].astype(str) + "%"
    
    # Figure should use numeric values
    fig_df = vil_caste.reset_index().melt(id_vars="village", value_vars=expected, var_name="Category", value_name="Count")
    # Calculate percentage for chart
    fig_df["Percentage"] = (fig_df["Count"] / fig_df.groupby("village")["Count"].transform("sum") * 100).round(2)
    
    fig2 = px.bar(fig_df.sort_values("Percentage", ascending=False), x="village", y="Percentage", color="Category", title="4.2. Caste wise number of families - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("4.2", "Caste wise number of families - Village level", table_df[cols_order], fig2)

def run_5_occupations(df):
    
    # 5.1 Panchayath Level
    gp_occ = df.groupby("primary_occupation").size().reset_index(name="No of households")
    gp_occ.rename(columns={"primary_occupation": "Primary Occupation"}, inplace=True)
    # FORMULA ALIGNMENT: Reference script uses sum of occupation counts (Share)
    share_total = gp_occ["No of households"].sum()
    gp_occ["Percentage"] = (gp_occ["No of households"] / share_total * 100).round(2) if share_total > 0 else 0
    gp_occ["No of households %"] = gp_occ["Percentage"].astype(str) + "%"
    
    fig_gp = px.bar(gp_occ.sort_values("Percentage", ascending=False), x="Primary Occupation", y="Percentage", title="5.1. Primary Occupations – Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("5.1", "Primary Occupations – Panchayat level", clean_df(gp_occ.drop(columns=["Percentage"])), fig_gp)

    # 5.2 Village Level
    vil_occ_pivot = df.pivot_table(index="primary_occupation", columns="village", aggfunc="size", fill_value=0).reset_index()
    vil_occ_pivot.rename(columns={"primary_occupation": "Primary Occupation"}, inplace=True)
    vil_occ_pivot["Grand Total"] = vil_occ_pivot.iloc[:, 1:].sum(axis=1)
    
    # Calculate percentage for bars
    vil_occ_melted = vil_occ_pivot.melt(id_vars="Primary Occupation", value_vars=vil_occ_pivot.columns[1:-1], var_name="village", value_name="Count")
    vil_total_per_vil = vil_occ_melted.groupby("village")["Count"].transform("sum")
    vil_occ_melted["Percentage"] = (vil_occ_melted["Count"] / vil_total_per_vil * 100).round(2)
    
    total_row_vals = vil_occ_pivot.iloc[:, 1:].sum().tolist()
    total_row = pd.DataFrame([["Grand Total"] + total_row_vals], columns=vil_occ_pivot.columns)
    vil_occ_table = pd.concat([vil_occ_pivot, total_row], ignore_index=True)
    
    fig_vil = px.bar(vil_occ_melted.sort_values("Percentage", ascending=False), x="Primary Occupation", y="Percentage", color="village", title="5.2. Primary Occupations – Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("5.2", "Primary Occupations – Village level", clean_df(vil_occ_table), fig_vil)

def run_6_income_sources(df):
    # Splitting multi-choice income sources using whitespace as per reference line 561
    df_split = df.assign(income_cleaned=df["income_sources"].astype(str).str.split()).explode("income_cleaned")
    df_split["income_cleaned"] = df_split["income_cleaned"].str.strip()

    # 6.1 Panchayath Level
    gp_inc = df_split.groupby("income_cleaned")["farmer_name"].nunique().reset_index(name="No of households")
    gp_inc.rename(columns={"income_cleaned": "Income Source"}, inplace=True)
    # FORMULA ALIGNMENT: Reference script uses sum of income source counts (Share)
    share_total = gp_inc["No of households"].sum()
    gp_inc["No of households %"] = (gp_inc["No of households"] / share_total * 100).round(2).astype(str) + "%"
    gp_inc["Percentage"] = (gp_inc["No of households"] / share_total * 100).round(2)
    
    fig_gp = px.bar(gp_inc.sort_values("Percentage", ascending=False), x="Income Source", y="Percentage", title="6.1. Income Sources – Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("6.1", "Income Sources – Panchayath level", clean_df(gp_inc.drop(columns=["Percentage"])), fig_gp)

    # 6.2 Village Level
    vil_inc_pivot = df_split.pivot_table(index="income_cleaned", columns="village", values="farmer_name", aggfunc="nunique", fill_value=0)
    vil_inc_pivot.columns = pd.MultiIndex.from_product([["No of households"], vil_inc_pivot.columns])
    vil_inc_pivot = vil_inc_pivot.reset_index()
    vil_inc_pivot.rename(columns={"income_cleaned": "Income Source"}, inplace=True)
    
    # Grand Total Row
    gt_vals = vil_inc_pivot.iloc[:, 1:].sum()
    gt_row = pd.DataFrame([["Grand Total"] + gt_vals.tolist()], columns=vil_inc_pivot.columns)
    vil_inc_table = pd.concat([vil_inc_pivot, gt_row], ignore_index=True)
    
    # Calculate percentage for bars
    vil_inc_melted = df_split.groupby(["village", "income_cleaned"])["farmer_name"].nunique().reset_index(name="Count")
    # FORMULA ALIGNMENT: Reference script uses sum of income source counts in village (Share)
    share_per_vil = vil_inc_melted.groupby("village")["Count"].transform("sum")
    vil_inc_melted["Percentage"] = (vil_inc_melted["Count"] / share_per_vil * 100).round(2)
    
    fig_vil = px.bar(vil_inc_melted.sort_values("Percentage", ascending=False), x="village", y="Percentage", color="income_cleaned", title="6.2. Income sources – Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("6.2", "Income sources – Village level", clean_df(vil_inc_table), fig_vil)

def run_7_job_cards(df):
    
    # 7.1 Panchayath Level
    df["job_card_status"] = normalize_yes_no(df["job_card_owned"])
    job_counts = df["job_card_status"].value_counts().reset_index()
    job_counts.columns = ["Job Card Status", "No of households"]
    total = job_counts["No of households"].sum()
    job_counts["No of households %"] = (job_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    job_counts["Job Card Status"] = job_counts["Job Card Status"].replace({"yes": "HHs having Job cards", "no": "HHs without Job cards"})
    
    fig_gp = px.pie(job_counts, values="No of households", names="Job Card Status", title="7.1. Job cards – Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("7.1", "Job cards – Panchayath level", job_counts, fig_gp)

    # 7.2 Village Level
    Y_L, N_L = "HHs having Job cards", "HHs without Job cards"
    vil_job = df.groupby(["village", "job_card_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_job: vil_job[c] = 0
    vil_job["Total"] = vil_job["yes"] + vil_job["no"]
    
    # Table Columns
    vil_job["Yes %"] = (vil_job["yes"] / vil_job["Total"] * 100).round(2)
    vil_job["No %"] = (vil_job["no"] / vil_job["Total"] * 100).round(2)
    
    # Grand Total Row
    gt_s = vil_job[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_job_table = pd.concat([vil_job, gt_row], ignore_index=True)
    vil_job_display = vil_job_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Legend Professional Fix: Use percentages for Y-axis
    vil_job["Job card HH %"] = vil_job["Yes %"]
    vil_job["No Job card HH %"] = (vil_job["no"] / vil_job["Total"] * 100).round(2)
    
    # Ensure display table doesn't have internal columns
    vil_job_display = vil_job_table.copy()
    vil_job_display["Yes %"] = vil_job_display["Yes %"].astype(str) + "%"
    vil_job_display["No %"] = vil_job_display["No %"].astype(str) + "%"
    vil_job_display.rename(columns={"yes": "HHs having Job cards", "no": "HHs without Job cards", "Yes %": "HHs having Job cards %", "No %": "HHs without Job cards %"}, inplace=True)
    
    fig_vil = px.bar(vil_job.sort_values("Yes %", ascending=False), x="village", y=["Job card HH %", "No Job card HH %"], title="7.2. Job cards – Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("7.2", "Job cards – Village level", clean_df(vil_job_display), fig_vil)

def run_8_head_of_households(df):
    # PREPROCESSING ALIGNMENT: Reference script lines 870-871
    df["is_women_headed_family"] = df["is_women_headed_family"].astype(str).str.lower().replace(["nan", "none"], "no")
    df["is_single_women_family"] = df["is_single_women_family"].astype(str).str.lower().replace(["nan", "none"], "no")
    
    total_hh = len(df)

    # 8.1 Head of the household - Panchayath level
    whf_count = (df["is_women_headed_family"] == "yes").sum()
    swf_count = (df["is_single_women_family"] == "yes").sum()
    
    gp_stats = pd.DataFrame([{
        "Women Headed Families": whf_count,
        "Single Women Families": swf_count,
        "Total": total_hh
    }])
    # FORMULA ALIGNMENT: Reference script lines 879-880
    gp_stats["Women Headed %"] = (whf_count / total_hh * 100).round(2).astype(str) + "%"
    gp_stats["Single Women %"] = (swf_count / total_hh * 100).round(2).astype(str) + "%"
    
    plot_df = pd.DataFrame({
        "Category": ["Women Headed", "Single Women"],
        "Percentage": [whf_count / total_hh * 100, swf_count / total_hh * 100]
    }).sort_values("Percentage", ascending=False)
    fig_gp = px.bar(plot_df, x="Category", y="Percentage", title="8.1. Head of the household - Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("8.1", "Head of the household - Panchayath level", gp_stats, fig_gp)

    # 8.2 Head of the household - village wise
    village_stats = df.groupby("village").agg(
        No_of_Women_Headed_Families=("is_women_headed_family", lambda x: (x == "yes").sum()),
        No_of_Single_Women_Families=("is_single_women_family", lambda x: (x == "yes").sum())
    ).reset_index()
    
    # FORMULA ALIGNMENT: Reference script lines 942-943 (Distribution logic)
    total_wh = village_stats["No_of_Women_Headed_Families"].sum()
    total_sw = village_stats["No_of_Single_Women_Families"].sum()
    
    village_stats["Women Headed Families %"] = (village_stats["No_of_Women_Headed_Families"] / total_wh * 100).round(2) if total_wh > 0 else 0
    village_stats["Single Women Families %"] = (village_stats["No_of_Single_Women_Families"] / total_sw * 100).round(2) if total_sw > 0 else 0
    
    # Grand Total Row
    total_vals = [village_stats["No_of_Women_Headed_Families"].sum(), village_stats["No_of_Single_Women_Families"].sum()]
    total_row = pd.DataFrame([[
        "Grand Total", total_vals[0], total_vals[1], 100.0, 100.0
    ]], columns=village_stats.columns)
    vil_head_table = pd.concat([village_stats, total_row], ignore_index=True)
    
    vil_head_display = vil_head_table.copy()
    vil_head_display["Women Headed Families %"] = vil_head_display["Women Headed Families %"].astype(str) + "%"
    vil_head_display["Single Women Families %"] = vil_head_display["Single Women Families %"].astype(str) + "%"
    vil_head_display.rename(columns={"No_of_Women_Headed_Families": "Women Headed", "No_of_Single_Women_Families": "Single Women"}, inplace=True)
    
    fig_vil = px.bar(village_stats.sort_values("Women Headed Families %", ascending=False), x="village", y=["Women Headed Families %", "Single Women Families %"], title="8.2. Head of the household - village wise", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("8.2", "Head of the household - village wise", clean_df(vil_head_display), fig_vil)

def run_9_migration(df):
    df["mig_status"] = normalize_yes_no(df["migration_from_family"])

    # 9.1 Migration at Panchayath level
    mig_counts = df["mig_status"].value_counts().reset_index()
    mig_counts.columns = ["Migration Status", "No of households"]
    total = mig_counts["No of households"].sum()
    mig_counts["No of households %"] = (mig_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    mig_counts["Migration Status"] = mig_counts["Migration Status"].replace({"yes": "Migrating families", "no": "Families without migration"})
    
    fig_gp = px.pie(mig_counts, values="No of households", names="Migration Status", title="9.1. Migration at Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("9.1", "Migration at Panchayath level", mig_counts, fig_gp)

    # 9.2 Migration at Village level
    Y_L, N_L = "Migrating families", "Families without migration"
    vil_mig = df.groupby(["village", "mig_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_mig: vil_mig[c] = 0
    vil_mig["Total"] = vil_mig["yes"] + vil_mig["no"]
    
    vil_mig["Yes %"] = (vil_mig["yes"] / vil_mig["Total"] * 100).round(2)
    vil_mig["No %"] = (vil_mig["no"] / vil_mig["Total"] * 100).round(2)
    
    gt_s = vil_mig[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_mig_table = pd.concat([vil_mig, gt_row], ignore_index=True)
    vil_mig_display = vil_mig_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Figure fix: Use percentage columns for Y-axis with descriptive labels
    vil_mig["Migrating families %"] = vil_mig["Yes %"]
    vil_mig["Families without migration %"] = vil_mig["No %"]
    fig_vil = px.bar(vil_mig.sort_values("Yes %", ascending=False), x="village", y=["Migrating families %", "Families without migration %"], title="9.2. Migration at Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("9.2", "Migration at Village level", clean_df(vil_mig_display), fig_vil)

    # 9.3 Caste wise Migration
    gp_caste_mig = df.groupby(["category", "mig_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in gp_caste_mig: gp_caste_mig[c] = 0
    gp_caste_mig["Total"] = gp_caste_mig["yes"] + gp_caste_mig["no"]
    gp_caste_mig["Yes %"] = (gp_caste_mig["yes"] / gp_caste_mig["Total"] * 100).round(2)
    
    gp_caste_mig["Migrating families %"] = gp_caste_mig["Yes %"]
    gp_caste_mig["Families without migration %"] = (gp_caste_mig["no"] / gp_caste_mig["Total"] * 100).round(2)
    gp_caste_mig_display = gp_caste_mig.rename(columns={"yes": "Migrating families", "no": "Families without migration"})
    fig_caste = px.bar(gp_caste_mig.sort_values("Yes %", ascending=False), x="category", y=["Migrating families %", "Families without migration %"], title="9.3. Caste wise migration – Panchayath level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("9.3", "Caste wise migration – Panchayath level", clean_df(gp_caste_mig_display), fig_caste)

def run_10_institutions(df):
    df["inst_mem"] = normalize_yes_no(df["member_in_institution"])

    # 10.1 Institution membership - Panchayath level
    mem_counts = df["inst_mem"].value_counts().reset_index()
    mem_counts.columns = ["Institution membership", "No of households"]
    total = mem_counts["No of households"].sum()
    mem_counts["No of households %"] = (mem_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    mem_counts["Institution membership"] = mem_counts["Institution membership"].replace({"yes": "HHs with Institution membership", "no": "HHs without Institution membership"})
    
    fig_gp = px.pie(mem_counts, values="No of households", names="Institution membership", title="10.1. Institution membership - Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("10.1", "Institution membership - Panchayath level", mem_counts, fig_gp)

    # 10.2 Institution membership - Village level
    Y_L, N_L = "HHs with Institution membership", "HHs without Institution membership"
    vil_mem = df.groupby(["village", "inst_mem"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_mem: vil_mem[c] = 0
    vil_mem["Total"] = vil_mem["yes"] + vil_mem["no"]
    
    # FORMULA ALIGNMENT: Use village-specific sum of yes/no
    vil_mem["Yes %"] = (vil_mem["yes"] / vil_mem["Total"] * 100).round(2)
    vil_mem["No %"] = (vil_mem["no"] / vil_mem["Total"] * 100).round(2)
    
    gt_s = vil_mem[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_mem_table = pd.concat([vil_mem, gt_row], ignore_index=True)
    vil_mem_display = vil_mem_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Legend Professional Fix: Use percentages for Y-axis
    vil_mem["Member HH %"] = vil_mem["Yes %"]
    vil_mem["Non-member HH %"] = (vil_mem["no"] / vil_mem["Total"] * 100).round(2)
    fig_vil = px.bar(vil_mem.sort_values("Yes %", ascending=False), x="village", y=["Member HH %", "Non-member HH %"], title="10.2. Institution membership - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("10.2", "Institution membership - Village level", clean_df(vil_mem_display), fig_vil)

def run_11_land_ownership(df):
    df["land_status"] = normalize_yes_no(df["land_owned"])

    # 11.1 Land ownership - Panchayat level
    land_counts = df["land_status"].value_counts().reset_index()
    land_counts.columns = ["Land Ownership Status", "No of households"]
    total = land_counts["No of households"].sum()
    land_counts["No of households %"] = (land_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    land_counts["Land Ownership Status"] = land_counts["Land Ownership Status"].replace({"yes": "HHs Land owning HHs", "no": "Landless Household"})
    
    fig_gp = px.pie(land_counts, values="No of households", names="Land Ownership Status", title="11.1. Land ownership - Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("11.1", "Land ownership - Panchayat level", land_counts, fig_gp)

    # 11.2 Caste wise land ownership
    gp_caste_land = df.groupby(["category", "land_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in gp_caste_land: gp_caste_land[c] = 0
    gp_caste_land["Total"] = gp_caste_land["yes"] + gp_caste_land["no"]
    gp_caste_land["Yes %"] = (gp_caste_land["yes"] / gp_caste_land["Total"] * 100).round(2)
    
    gp_caste_land["HHs landowning %"] = gp_caste_land["Yes %"]
    gp_caste_land["landless households %"] = (gp_caste_land["no"] / gp_caste_land["Total"] * 100).round(2)
    gp_caste_land_display = gp_caste_land.rename(columns={"yes": "HHs landowning", "no": "landless households"})
    gp_caste_land_display["Land owning %"] = gp_caste_land["Yes %"].astype(str) + "%"
    
    # Chart with percentages
    fig_caste = px.bar(gp_caste_land.sort_values("Yes %", ascending=False), x="category", y=["HHs landowning %", "landless households %"], title="11.2. Caste wise land ownership - Panchayat level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("11.2", "Caste wise land ownership - Panchayat level", clean_df(gp_caste_land_display), fig_caste)

    # 11.3 Land ownership - Village level
    Y_L, N_L = "HHs Land owning HHs", "Landless Household"
    vil_land = df.groupby(["village", "land_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_land: vil_land[c] = 0
    vil_land["Total"] = vil_land["yes"] + vil_land["no"]
    
    vil_land["Yes %"] = (vil_land["yes"] / vil_land["Total"] * 100).round(2)
    vil_land["No %"] = (vil_land["no"] / vil_land["Total"] * 100).round(2)
    
    gt_s = vil_land[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_land_table = pd.concat([vil_land, gt_row], ignore_index=True)
    vil_land_display = vil_land_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Figure fix: Use percentages for Y-axis
    vil_land["HHs landowning %"] = vil_land["Yes %"]
    vil_land["landless households %"] = (vil_land["no"] / vil_land["Total"] * 100).round(2)
    fig_vil = px.bar(vil_land.sort_values("Yes %", ascending=False), x="village", y=["HHs landowning %", "landless households %"], title="11.3. Land ownership - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("11.3", "Land ownership - Village level", clean_df(vil_land_display), fig_vil)

def run_12_land_leased(df):
    df["land_leased_in"] = normalize_yes_no(df["land_leased_in"])
    df["land_leased_out"] = normalize_yes_no(df["land_leased_out"])

    # No deduplication as per reference
    total_hh = len(df)

    # 12.1 Panchayath Level
    li_count = (df["land_leased_in"] == "yes").sum()
    lo_count = (df["land_leased_out"] == "yes").sum()
    
    gp_lease = pd.DataFrame([{
        "Leased In": li_count,
        "Leased Out": lo_count,
        "Total Households": total_hh
    }])
    # FORMULA ALIGNMENT: Prevalence within GP
    gp_lease["Leased In %"] = (li_count / total_hh * 100).round(2).astype(str) + "%"
    gp_lease["Leased Out %"] = (lo_count / total_hh * 100).round(2).astype(str) + "%"
    
    # Chart (Descending)
    plot_df = pd.DataFrame({
        "Category": ["Leased In", "Leased Out"],
        "Percentage": [li_count / total_hh * 100, lo_count / total_hh * 100]
    }).sort_values("Percentage", ascending=False)
    fig_gp = px.bar(plot_df, x="Category", y="Percentage", title="12.1. Panchayath Level Land leased in and leased out", color_discrete_sequence=COLORS)
    save_and_render("12.1", "Panchayath Level Land leased in and leased out", clean_df(gp_lease), fig_gp)

    # 12.2 Village Level
    vil_lease = df.groupby("village").agg(
        Leased_In=("land_leased_in", lambda x: (x == "yes").sum()),
        Leased_Out=("land_leased_out", lambda x: (x == "yes").sum()),
        Total_HH=("farmer_name", "count")
    ).reset_index()
    
    vil_lease["Leased In %"] = (vil_lease["Leased_In"] / vil_lease["Total_HH"] * 100).round(2)
    vil_lease["Leased Out %"] = (vil_lease["Leased_Out"] / vil_lease["Total_HH"] * 100).round(2)
    
    total_vals = [vil_lease["Leased_In"].sum(), vil_lease["Leased_Out"].sum(), vil_lease["Total_HH"].sum()]
    total_row = pd.DataFrame([[
        "Grand Total", total_vals[0], total_vals[1], total_vals[2],
        round(total_vals[0]/total_vals[2]*100, 2) if total_vals[2] > 0 else 0,
        round(total_vals[1]/total_vals[2]*100, 2) if total_vals[2] > 0 else 0
    ]], columns=vil_lease.columns)
    vil_lease_table = pd.concat([vil_lease, total_row], ignore_index=True)

    # Final formatting for table
    vil_lease_table_formatted = vil_lease_table.copy()
    vil_lease_table_formatted["Leased In %"] = vil_lease_table_formatted["Leased In %"].astype(str) + "%"
    vil_lease_table_formatted["Leased Out %"] = vil_lease_table_formatted["Leased Out %"].astype(str) + "%"
    vil_lease_table_formatted.rename(columns={"Leased_In": "Leased In", "Leased_Out": "Leased Out", "Total_HH": "Total Households"}, inplace=True)
    
    fig_vil = px.bar(vil_lease.sort_values("Leased In %", ascending=False), x="village", y=["Leased In %", "Leased Out %"], title="12.2. Village wise Land leased in and leased out", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("12.2", "Village wise Land leased in and leased out", clean_df(vil_lease_table_formatted), fig_vil)

def run_13_livestock(df):
    df["live_status"] = normalize_yes_no(df["livestock_owned"])

    # 13.1 Livestock-owning households - Panchayat level
    live_counts = df["live_status"].value_counts().reset_index()
    live_counts.columns = ["Livestock Status", "No of households"]
    total = live_counts["No of households"].sum()
    live_counts["No of households %"] = (live_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    live_counts["Livestock Status"] = live_counts["Livestock Status"].replace({"yes": "Livestock owned Households", "no": "Households without livestock"})
    
    fig_gp = px.pie(live_counts, values="No of households", names="Livestock Status", title="13.1. Livestock-owning households - Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("13.1", "Livestock-owning households - Panchayat level", live_counts, fig_gp)

    # 13.2 Livestock owning households - Village level
    Y_L, N_L = "Livestock owned Households", "Households without livestock"
    vil_live = df.groupby(["village", "live_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_live: vil_live[c] = 0
    vil_live["Total"] = vil_live["yes"] + vil_live["no"]
    
    vil_live["Yes %"] = (vil_live["yes"] / vil_live["Total"] * 100).round(2)
    vil_live["No %"] = (vil_live["no"] / vil_live["Total"] * 100).round(2)
    
    gt_s = vil_live[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_live_table = pd.concat([vil_live, gt_row], ignore_index=True)
    vil_live_display = vil_live_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Figure fix: Use percentages for Y-axis
    vil_live["Livestock owning HH %"] = vil_live["Yes %"]
    vil_live["Non-livestock HH %"] = (vil_live["no"] / vil_live["Total"] * 100).round(2)
    fig_vil = px.bar(vil_live.sort_values("Yes %", ascending=False), x="village", y=["Livestock owning HH %", "Non-livestock HH %"], title="13.2. Livestock owning households - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("13.2", "Livestock owning households - Village level", clean_df(vil_live_display), fig_vil)

def run_14_soil_types(df_plot):
    # 14.1 Soil types and extent Panchayat level
    gp_soil = df_plot.groupby("soil_type")["Area"].sum().reset_index()
    gp_soil.columns = ["Soil Type", "Total Extent (Acres)"]
    gp_soil_table = gp_soil.copy()
    gp_total = gp_soil["Total Extent (Acres)"].sum()
    gp_soil_table["Extent%"] = (gp_soil["Total Extent (Acres)"] / gp_total * 100).round(2).astype(str) + "%"
    
    # Use numeric for chart and sort descending
    gp_soil["Extent%"] = (gp_soil["Total Extent (Acres)"] / gp_total * 100).round(2)
    fig_gp = px.pie(gp_soil, values="Total Extent (Acres)", names="Soil Type", title="14.1. Soil types and extent Panchayat level", hole=0.5, color_discrete_sequence=COLORS)
    save_and_render("14.1", "Soil types and extent Panchayat level", gp_soil_table, fig_gp)

    # 14.2 Soil types and extent Village level
    vil_soil = df_plot.pivot_table(index="soil_type", columns="village", values="Area", aggfunc="sum", fill_value=0).reset_index()
    vil_soil.rename(columns={"soil_type": "Soil Type"}, inplace=True)
    vil_soil["Total Extent"] = vil_soil.iloc[:, 1:].sum(axis=1)
    
    # Grand Total Row
    gt_vals = vil_soil.iloc[:, 1:].sum()
    gt_row = pd.DataFrame([["Grand Total"] + gt_vals.tolist()], columns=vil_soil.columns)
    vil_soil_table = pd.concat([vil_soil, gt_row], ignore_index=True)
    
    # Chart
    vil_soil_melted = df_plot.groupby(["village", "soil_type"])["Area"].sum().reset_index()
    vil_soil_total = vil_soil_melted.groupby("village")["Area"].transform("sum")
    vil_soil_melted["Extent%"] = (vil_soil_melted["Area"] / vil_soil_total * 100).round(2)
    fig_vil = px.bar(vil_soil_melted.sort_values("Extent%", ascending=False), x="village", y="Extent%", color="soil_type", title="14.2. Soil types and extent Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("14.2", "Soil types and extent Village level", clean_df(vil_soil_table), fig_vil)

def run_15_irrigation(df_plot, df_hh):
    
    # 15.1 & 15.2: Household Access (Yes/No)
    hh_irr_col = "Irrigation_sources" if "Irrigation_sources" in df_hh.columns else "is_any_irrigation_source"
    df_hh["irr_access"] = normalize_yes_no(df_hh[hh_irr_col])

    # 15.1 Panchayath level
    irr_counts = df_hh["irr_access"].value_counts().reset_index()
    irr_counts.columns = ["Irrigation Access", "No of households"]
    total = irr_counts["No of households"].sum()
    irr_counts["No of households %"] = (irr_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    irr_counts["Irrigation Access"] = irr_counts["Irrigation Access"].replace({"yes": "Households having irrigation sources", "no": "Households without irrigation sources"})
    
    fig_gp = px.pie(irr_counts, values="No of households", names="Irrigation Access", title="15.1. Households with Irrigation sources - Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("15.1", "Households with Irrigation sources - Panchayath level", irr_counts, fig_gp)

    # 15.2 Village level
    Y_L, N_L = "Households having irrigation sources", "Households without irrigation sources"
    vil_irr = df_hh.groupby(["village", "irr_access"]).size().unstack(fill_value=0).reset_index()
    for col in ["yes", "no"]:
        if col not in vil_irr: vil_irr[col] = 0
    vil_irr["Total"] = vil_irr["yes"] + vil_irr["no"]
    
    vil_irr["Yes %"] = (vil_irr["yes"] / vil_irr["Total"] * 100).round(2)
    vil_irr["No %"] = (vil_irr["no"] / vil_irr["Total"] * 100).round(2)
    
    gt_s = vil_irr[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_irr_table = pd.concat([vil_irr, gt_row], ignore_index=True)
    vil_irr_display = vil_irr_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Figure fix: Use percentages for Y-axis
    vil_irr["Irrigated HH %"] = vil_irr["Yes %"]
    vil_irr["Non-irrigated HH %"] = (vil_irr["no"] / vil_irr["Total"] * 100).round(2)
    fig_vil = px.bar(vil_irr.sort_values("Yes %", ascending=False), x="village", y=["Irrigated HH %", "Non-irrigated HH %"], title="15.2. Households with Irrigation sources - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("15.2", "Households with Irrigation sources - Village level", clean_df(vil_irr_display), fig_vil)

    # 15.3 & 15.4: Type of Irrigation Sources (Exploded multi-select)
    type_col = "type_Irrigation_sources" if "type_Irrigation_sources" in df_hh.columns else "Irrigation_sources"
    df_hh_clean = df_hh[df_hh[type_col].notna()].copy()
    df_hh_clean[type_col] = df_hh_clean[type_col].astype(str).str.split()
    df_exploded = df_hh_clean.explode(type_col)
    df_exploded[type_col] = df_exploded[type_col].str.strip().str.lower()
    # Filter out 'yes'/'no' if they were mixed in
    df_exploded = df_exploded[~df_exploded[type_col].isin(["yes", "no", "nan", "none"])]
    df_exploded[type_col] = df_exploded[type_col].apply(lambda x: str(x).replace('_', ' ').title() if isinstance(x, str) and '_' in str(x) else x)

    # 15.3 Panchayath Level Types
    gp_type = df_exploded.groupby(type_col)["farmer_name"].nunique().reset_index(name="No of households")
    gp_total_hh = df_hh["farmer_name"].nunique()
    gp_type["No of households %"] = (gp_type["No of households"] / gp_total_hh * 100).round(2)
    gp_type["Percentage_Str"] = gp_type["No of households %"].astype(str) + "%"
    gp_type.rename(columns={type_col: "Irrigation Type"}, inplace=True)
    
    fig_gp_type = px.bar(gp_type.sort_values("No of households %", ascending=False), x="Irrigation Type", y="No of households %", title="15.3. Type of irrigation sources - Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("15.3", "Type of irrigation sources - Panchayat level", clean_df(gp_type.drop(columns=["No of households %"]).rename(columns={"Percentage_Str": "No of households %"})), fig_gp_type)

    # 15.4 Village Level Types
    vil_type = df_exploded.pivot_table(index=type_col, columns="village", values="farmer_name", aggfunc="nunique", fill_value=0).reset_index()
    vil_type.columns.name = "No of household"
    vil_type.rename(columns={type_col: "Irrigation Type"}, inplace=True)
    vil_type["Grand Total"] = vil_type.iloc[:, 1:].sum(axis=1)
    
    # Calculate percentage for bars - Denominator: Total farmers who have irrigation (as per reference 15.4)
    vil_type_melted = vil_type.melt(id_vars="Irrigation Type", value_vars=vil_type.columns[1:-1], var_name="village", value_name="Count")
    vil_total_irr = vil_type_melted.groupby("village")["Count"].transform("sum")
    vil_type_melted["Percentage"] = (vil_type_melted["Count"] / vil_total_irr * 100).round(2)
    
    fig_vil_type = px.bar(vil_type_melted.sort_values("Percentage", ascending=False), x="village", y="Percentage", color="Irrigation Type", barmode="group", title="15.4. Type of irrigation sources - Village level", color_discrete_sequence=COLORS)
    save_and_render("15.4", "Type of irrigation sources - Village level", clean_df(vil_type), fig_vil_type)

    # 15.5 & 15.6: Caste-wise Irrigation Sources
    # 15.5 Panchayath level
    gp_caste_irr = df_exploded.pivot_table(index="category", columns=type_col, values="farmer_name", aggfunc="nunique", fill_value=0).reset_index()
    gp_caste_irr_melted = df_exploded.groupby(["category", type_col])["farmer_name"].nunique().reset_index(name="Count")
    gp_caste_total = gp_caste_irr_melted.groupby("category")["Count"].transform("sum")
    gp_caste_irr_melted["Percentage"] = (gp_caste_irr_melted["Count"] / gp_caste_total * 100).round(2)
    fig_gp_caste = px.bar(gp_caste_irr_melted.sort_values("Percentage", ascending=False), x=type_col, y="Percentage", color="category", barmode="group", color_discrete_sequence=COLORS)
    fig_gp_caste.update_layout(title="15.5. Irrigation sources by category - panchayat level")
    save_and_render("15.5", "Irrigation sources by category - panchayat level", clean_df(gp_caste_irr), fig_gp_caste)

    # 15.6 Village level
    vil_caste_irr = df_exploded.pivot_table(index=type_col, columns=["village", "category"], values="farmer_name", aggfunc="nunique", fill_value=0).reset_index()
    vil_caste_irr_melted = df_exploded.groupby(["village", "category", type_col])["farmer_name"].nunique().reset_index(name="Count")
    vil_caste_total = vil_caste_irr_melted.groupby(["village", "category"])["Count"].transform("sum")
    vil_caste_irr_melted["Percentage"] = (vil_caste_irr_melted["Count"] / vil_caste_total * 100).round(2)
    fig_vil_caste = px.bar(vil_caste_irr_melted.sort_values("Percentage", ascending=False), x="village", y="Percentage", color=type_col, facet_col="category", barmode="group", color_discrete_sequence=COLORS)
    fig_vil_caste.update_layout(title="15.6. Irrigation sources by category - village level")
    save_and_render("15.6", "Irrigation sources by category - village level", clean_df(vil_caste_irr), fig_vil_caste)

    # 15.7 & 15.8: Extent of Irrigation (Area-wise from Plot Data)
    plot_irr_col = "Irrigation_sources" if "Irrigation_sources" in df_plot.columns else "type_Irrigation_sources"
    # 15.7 Panchayath Level Extent
    gp_area_irr = df_plot.groupby(plot_irr_col)["Area"].sum().reset_index()
    gp_total_area = gp_area_irr["Area"].sum()
    gp_area_irr["Extent%"] = (gp_area_irr["Area"] / gp_total_area * 100).round(2)
    gp_area_irr["Extent%_Str"] = gp_area_irr["Extent%"].astype(str) + "%"
    gp_area_irr.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp_area = px.bar(gp_area_irr.sort_values("Extent%", ascending=False), x=plot_irr_col, y="Extent%", title="15.7. Extent of irrigation at Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("15.7", "Extent of irrigation at Panchayath level", clean_df(gp_area_irr.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp_area)

    # 15.8 Village Level Extent
    vil_area_irr = df_plot.groupby(["village", plot_irr_col])["Area"].sum().reset_index()
    vil_total_area = vil_area_irr.groupby("village")["Area"].transform("sum")
    vil_area_irr["Extent%"] = (vil_area_irr["Area"] / vil_total_area * 100).round(2)
    vil_area_irr.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil_area = px.bar(vil_area_irr, x="village", y="Extent%", color=plot_irr_col, barmode="group", title="15.8. Extent of irrigation at village level", color_discrete_sequence=COLORS)
    save_and_render("15.8", "Extent of irrigation at village level", clean_df(vil_area_irr), fig_vil_area)

def run_16_lulc(df_plot):
    # 16.1 Panchayat Level LULC
    gp_lulc = df_plot.groupby("LULC")["Area"].sum().reset_index()
    gp_lulc.columns = ["LULC", "Total Extent (Acres)"]
    gp_lulc_table = gp_lulc.copy()
    gp_lulc_table["Extent%"] = (gp_lulc["Total Extent (Acres)"] / gp_lulc["Total Extent (Acres)"].sum() * 100).round(2).astype(str) + "%"
    fig_gp = px.pie(gp_lulc, values="Total Extent (Acres)", names="LULC", title="16.1. LULC WISE EXTENT - Panchayat Level", hole=0.5, color_discrete_sequence=COLORS)
    save_and_render("16.1", "LULC Distribution - Panchayat Level", gp_lulc_table, fig_gp)

    # 16.2 Village Level LULC
    vil_lulc = df_plot.pivot_table(index="LULC", columns="village", values="Area", aggfunc="sum", fill_value=0).reset_index()
    vil_lulc.rename(columns={"LULC": "LULC Type"}, inplace=True)
    vil_lulc["Total Extent"] = vil_lulc.iloc[:, 1:].sum(axis=1)
    
    # Grand Total Row
    gt_vals = vil_lulc.iloc[:, 1:].sum()
    gt_row = pd.DataFrame([["Grand Total"] + gt_vals.tolist()], columns=vil_lulc.columns)
    vil_lulc_table = pd.concat([vil_lulc, gt_row], ignore_index=True)
    
    # Calculate percentage for bars
    vil_lulc_melted = df_plot.groupby(["village", "LULC"])["Area"].sum().reset_index()
    vil_total_area = vil_lulc_melted.groupby("village")["Area"].transform("sum")
    vil_lulc_melted["Extent %"] = (vil_lulc_melted["Area"] / vil_total_area * 100).round(2)
    fig_vil = px.bar(vil_lulc_melted.sort_values("Extent %", ascending=False), x="village", y="Extent %", color="LULC", barmode="group", title="16.2. LULC distribution - Village Level", color_discrete_sequence=COLORS)
    save_and_render("16.2", "LULC distribution - Village Level", clean_df(vil_lulc_table), fig_vil)

def run_17_cropping_systems(df_plot):
    # 17.2 Cropping system wise no of households - Panchayat
    gp_cs = df_plot.groupby("cropping_system")["farmer_name"].nunique().reset_index(name="No of households").sort_values("No of households", ascending=False)
    gp_cs_table = gp_cs.copy()
    gp_cs_table["Households %"] = (gp_cs["No of households"] / gp_cs["No of households"].sum() * 100).round(2).astype(str) + "%"
    fig_gp = px.pie(gp_cs, values="No of households", names="cropping_system", title="17.2. Cropping system wise no of households - Panchayat Level", color_discrete_sequence=COLORS)
    save_and_render("17.2", "Cropping system wise no of households - Panchayat Level", clean_df(gp_cs_table), fig_gp)

    # 17.3 Village wise no of households
    vil_cs = df_plot.groupby(["village", "cropping_system"])["farmer_name"].nunique().reset_index(name="No of households").sort_values("No of households", ascending=False)
    # Pivot for table view
    vil_cs_pivot = vil_cs.pivot_table(index="village", columns="cropping_system", values="No of households", fill_value=0).reset_index()
    crops = [c for c in vil_cs_pivot.columns if c != "village"]
    vil_cs_pivot["Total"] = vil_cs_pivot[crops].sum(axis=1)
    
    # Interleave Percentages
    table_df = vil_cs_pivot.copy()
    cols_order = ["village"]
    for c in crops:
        table_df[f"{c} %"] = (table_df[c] / table_df["Total"] * 100).round(2).astype(str) + "%"
        cols_order.extend([c, f"{c} %"])
    cols_order.append("Total")
    
    # Grand Total Row
    gt_s = vil_cs_pivot[crops + ["Total"]].sum()
    gt_row = {"village": "Grand Total"}
    for c in crops:
        gt_row[c] = gt_s[c]
        gt_row[f"{c} %"] = f"{(gt_s[c]/gt_s['Total']*100):.2f}%" if gt_s["Total"] > 0 else "0.00%"
    gt_row["Total"] = gt_s["Total"]
    table_df = pd.concat([table_df, pd.DataFrame([gt_row])], ignore_index=True)
    
    # Calculate percentage for bars
    vil_cs_plot = vil_cs.copy()
    vil_cs_total = vil_cs_plot.groupby("village")["No of households"].transform("sum")
    vil_cs_plot["Percentage"] = (vil_cs_plot["No of households"] / vil_cs_total * 100).round(2)
    
    fig_vil = px.bar(vil_cs_plot.sort_values("Percentage", ascending=False), x="village", y="Percentage", color="cropping_system", title="17.3. Cropping system wise no of farmers -Village wise", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("17.3", "Cropping system wise no of farmers -Village wise", clean_df(table_df[cols_order]), fig_vil)

    # 17.4(a) Extent of Irrigated and Rainfed Area - Panchayath Level
    gp_extent = df_plot.groupby("cropping_system")["Area"].sum().reset_index()
    total_area = gp_extent["Area"].sum()
    gp_extent["Extent%"] = (gp_extent["Area"] / total_area * 100).round(2)
    gp_extent.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp_extent = px.bar(gp_extent, x="cropping_system", y="Extent%", text="Extent%", title="17.4(a) Irrigated and Rainfed Area - Panchayath Level", color="cropping_system", color_discrete_sequence=COLORS)
    save_and_render("17.4(a)", "Irrigated and Rainfed Area - Panchayath Level", clean_df(gp_extent), fig_gp_extent)

    # 17.4(b) Village-wise Irrigated vs Rainfed Area
    vil_extent = df_plot.groupby(["village", "cropping_system"])["Area"].sum().reset_index()
    vil_total_area = vil_extent.groupby("village")["Area"].transform("sum")
    vil_extent["Extent%"] = (vil_extent["Area"] / vil_total_area * 100).round(2)
    vil_extent.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil_extent = px.bar(vil_extent, x="village", y="Extent%", color="cropping_system", barmode="group", text="Extent%", title="17.4(b) Village-wise Irrigated vs Rainfed Area", color_discrete_sequence=COLORS)
    save_and_render("17.4(b)", "Village-wise Irrigated vs Rainfed Area", clean_df(vil_extent), fig_vil_extent)

def run_18_cropping_methods(df_plot):
    method_col = next((c for c in ["croping_method", "cropping_method", "cropping method"] if c in df_plot.columns), "cropping_method")
    practice_col = next((c for c in ["cropping_practice", "cropping practice"] if c in df_plot.columns), "cropping_practice")
    crop_col = next((c for c in ["crop_names", "crop_name", "crop names"] if c in df_plot.columns), "crop_name")
    season_col = next((c for c in ["cropping_seasons", "season", "cropping_season", "season_name"] if c in df_plot.columns), "season")

    # 18.2 No of households
    gp_cm = df_plot.groupby(method_col)["farmer_name"].nunique().reset_index(name="No of households")
    gp_total = gp_cm["No of households"].sum()
    gp_cm["Households %"] = (gp_cm["No of households"] / gp_total * 100).round(2).astype(str) + "%"
    fig_gp_cm = px.bar(gp_cm, x=method_col, y="No of households", title="18.2. Cropping method wise no of farmers", color_discrete_sequence=COLORS)
    save_and_render("18.2", "Cropping method wise no of farmers", clean_df(gp_cm), fig_gp_cm)

    vil_cm = df_plot.groupby(["village", method_col])["farmer_name"].nunique().reset_index(name="No of households")
    vil_total = vil_cm.groupby("village")["No of households"].transform("sum")
    vil_cm["Households %"] = (vil_cm["No of households"] / vil_total * 100).round(2)
    fig_vil_cm = px.bar(vil_cm.sort_values("Households %", ascending=False), x="village", y="Households %", color=method_col, barmode="group", text="Households %", title="18.2(b) Cropping method wise no of households - Village level", color_discrete_sequence=COLORS)
    save_and_render("18.2(b)", "Cropping method wise no of households - Village level", clean_df(vil_cm), fig_vil_cm)

    # 18.3 & 18.4 Cropping method wise extent
    gp_me = df_plot.groupby(method_col)["Area"].sum().reset_index()
    gp_total_area = gp_me["Area"].sum()
    gp_me["Extent%"] = (gp_me["Area"] / gp_total_area * 100).round(2)
    gp_me["Extent%_Str"] = gp_me["Extent%"].astype(str) + "%"
    gp_me.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp_me = px.bar(gp_me.sort_values("Extent%", ascending=False), x=method_col, y="Extent%", title="18.3. Cropping method wise extent - Panchayat Level", color_discrete_sequence=COLORS)
    save_and_render("18.3", "Cropping method wise extent - Panchayat Level", clean_df(gp_me.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp_me)

    vil_me = df_plot.groupby(["village", method_col])["Area"].sum().reset_index()
    vil_total_area = vil_me.groupby("village")["Area"].transform("sum")
    vil_me["Extent%"] = (vil_me["Area"] / vil_total_area * 100).round(2)
    vil_me.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil_me = px.bar(vil_me.sort_values("Extent%", ascending=False), x="village", y="Extent%", color=method_col, barmode="group", title="18.4. Cropping method wise extent - Village level", color_discrete_sequence=COLORS)
    save_and_render("18.4", "Cropping method wise extent - Village level", clean_df(vil_me), fig_vil_me)

    # 18.5 Cropping practice wise extent
    gp_pe = df_plot.groupby(practice_col)["Area"].sum().reset_index()
    gp_total_area = gp_pe["Area"].sum()
    gp_pe["Extent%"] = (gp_pe["Area"] / gp_total_area * 100).round(2)
    gp_pe["Extent%_Str"] = gp_pe["Extent%"].astype(str) + "%"
    gp_pe.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp_pe = px.bar(gp_pe.sort_values("Extent%", ascending=False), x=practice_col, y="Extent%", title="18.5(a) Cropping practice wise extent - Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("18.5(a)", "Cropping practice wise extent - Panchayat level", clean_df(gp_pe.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp_pe)

    vil_pe = df_plot.groupby(["village", practice_col])["Area"].sum().reset_index()
    vil_total_area = vil_pe.groupby("village")["Area"].transform("sum")
    vil_pe["Extent%"] = (vil_pe["Area"] / vil_total_area * 100).round(2)
    vil_pe.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil_pe = px.bar(vil_pe.sort_values("Extent%", ascending=False), x="village", y="Extent%", color=practice_col, barmode="group", title="18.5(b) Cropping practice wise extent - Village level", color_discrete_sequence=COLORS)
    save_and_render("18.5(b)", "Cropping practice wise extent - Village level", clean_df(vil_pe), fig_vil_pe)

    # 18.6 Crops wise extent
    gp_ce = df_plot.groupby(crop_col)["Area"].sum().reset_index()
    gp_total_area = gp_ce["Area"].sum()
    gp_ce["Extent%"] = (gp_ce["Area"] / gp_total_area * 100).round(2)
    gp_ce["Extent%_Str"] = gp_ce["Extent%"].astype(str) + "%"
    gp_ce.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp_ce = px.bar(gp_ce.sort_values("Extent%", ascending=False), x=crop_col, y="Extent%", title="18.6(a) Crops wise extent - Panchayat level", color_discrete_sequence=COLORS)
    save_and_render("18.6(a)", "Crops wise extent - Panchayat level", clean_df(gp_ce.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp_ce)

    vil_ce = df_plot.groupby(["village", crop_col])["Area"].sum().reset_index()
    vil_total_area = vil_ce.groupby("village")["Area"].transform("sum")
    vil_ce["Extent%"] = (vil_ce["Area"] / vil_total_area * 100).round(2)
    vil_ce.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil_ce = px.bar(vil_ce.sort_values("Extent%", ascending=False), x="village", y="Extent%", color=crop_col, barmode="group", title="18.6(b) Crops wise extent - Village level", color_discrete_sequence=COLORS)
    save_and_render("18.6(b)", "Crops wise extent - Village level", clean_df(vil_ce), fig_vil_ce)

    # 18.7 Cropping seasons wise extent
    if season_col in df_plot.columns:
        gp_se = df_plot.groupby(season_col)["Area"].sum().reset_index()
        gp_total_area = gp_se["Area"].sum()
        gp_se["Extent%"] = (gp_se["Area"] / gp_total_area * 100).round(2)
        gp_se["Extent%_Str"] = gp_se["Extent%"].astype(str) + "%"
        gp_se.rename(columns={"Area": "Extent (ac)"}, inplace=True)
        fig_gp_se = px.bar(gp_se.sort_values("Extent%", ascending=False), x=season_col, y="Extent%", title="18.7(a) Cropping seasons and extent - Panchayat level", color_discrete_sequence=COLORS)
        save_and_render("18.7(a)", "Cropping seasons and extent - Panchayat level", clean_df(gp_se.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp_se)
    
        vil_se = df_plot.groupby(["village", season_col])["Area"].sum().reset_index()
        vil_total_area = vil_se.groupby("village")["Area"].transform("sum")
        vil_se["Extent%"] = (vil_se["Area"] / vil_total_area * 100).round(2)
        vil_se.rename(columns={"Area": "Extent (ac)"}, inplace=True)
        fig_vil_se = px.bar(vil_se, x="village", y="Extent%", color=season_col, barmode="group", title="18.7(b) Cropping seasons and extent - Village level", color_discrete_sequence=COLORS)
        save_and_render("18.7(b)", "Cropping seasons and extent - Village level", clean_df(vil_se), fig_vil_se)

def run_19_fallow_lands(df_plot):
    if "from_how_many_years" not in df_plot.columns:
        return

    # 19.1 Panchayat Level Extent
    gp_fallow = df_plot[df_plot["from_how_many_years"].notna()].groupby("from_how_many_years")["Area"].sum().reset_index()
    gp_total = gp_fallow["Area"].sum()
    gp_fallow["Extent%"] = (gp_fallow["Area"] / gp_total * 100).round(2)
    gp_fallow["Extent%_Str"] = gp_fallow["Extent%"].astype(str) + "%"
    gp_fallow.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_gp = px.bar(gp_fallow.sort_values("Extent%", ascending=False), x="from_how_many_years", y="Extent%", title="19.1. Fallow Lands Extent - Panchayat Level", color_discrete_sequence=COLORS)
    save_and_render("19.1", "Fallow Lands Extent - Panchayat Level", clean_df(gp_fallow.drop(columns=["Extent%"]).rename(columns={"Extent%_Str": "Extent%"})), fig_gp)

    # 19.2 Fallow land types households
    gp_hh_fallow = df_plot[df_plot["from_how_many_years"].notna()].groupby("from_how_many_years")["farmer_name"].nunique().reset_index(name="No of households")
    gp_hh_total = df_plot["farmer_name"].nunique()
    gp_hh_fallow["Households %"] = (gp_hh_fallow["No of households"] / gp_hh_total * 100).round(2)
    gp_hh_fallow["Households %_Str"] = gp_hh_fallow["Households %"].astype(str) + "%"
    
    fig_gp_hh = px.bar(gp_hh_fallow.sort_values("Households %", ascending=False), x="from_how_many_years", y="Households %", title="19.2. Fallow land types households - Panchayath level", color_discrete_sequence=COLORS)
    save_and_render("19.2", "Fallow land types households - Panchayath level", clean_df(gp_hh_fallow.drop(columns=["Households %"]).rename(columns={"Households %_Str": "Households %"})), fig_gp_hh)

    # 19.3 Fallow land - extent of fallow (Village)
    vil_fallow = df_plot[df_plot["from_how_many_years"].notna()].groupby(["village", "from_how_many_years"])["Area"].sum().reset_index()
    vil_total = vil_fallow.groupby("village")["Area"].transform("sum")
    vil_fallow["Extent%"] = (vil_fallow["Area"] / vil_total * 100).round(2)
    vil_fallow.rename(columns={"Area": "Extent (ac)"}, inplace=True)
    fig_vil = px.bar(vil_fallow.sort_values("Extent%", ascending=False), x="village", y="Extent%", color="from_how_many_years", barmode="group", text="Extent%", title="19.3. Fallow land - extent of fallow - Village level", color_discrete_sequence=COLORS)
    save_and_render("19.3", "Fallow land - extent of fallow - Village level", clean_df(vil_fallow), fig_vil)

def run_20_kitchen_gardens(df):
    df["kg_status"] = normalize_yes_no(df["hh_has_kitchen_garden"])

    # 20.1 Panchayath Level
    kg_counts = df["kg_status"].value_counts().reset_index()
    kg_counts.columns = ["Kitchen Garden Status", "No of households"]
    total = kg_counts["No of households"].sum()
    kg_counts["No of households %"] = (kg_counts["No of households"] / total * 100).round(2).astype(str) + "%"
    kg_counts["Kitchen Garden Status"] = kg_counts["Kitchen Garden Status"].replace({"yes": "HH with kitchen garden", "no": "HH without kitchen garden"})
    
    fig_gp = px.pie(kg_counts, values="No of households", names="Kitchen Garden Status", title="20. KITCHEN GARDENS", color_discrete_sequence=COLORS)
    save_and_render("20.1", "Kitchen Gardens (Panchayath)", kg_counts, fig_gp)

    Y_L, N_L = "HH with kitchen garden", "HH without kitchen garden"
    vil_kg = df.groupby(["village", "kg_status"]).size().unstack(fill_value=0).reset_index()
    for c in ["yes", "no"]:
        if c not in vil_kg: vil_kg[c] = 0
    vil_kg["Total"] = vil_kg["yes"] + vil_kg["no"]
    
    vil_kg["Yes %"] = (vil_kg["yes"] / vil_kg["Total"] * 100).round(2)
    vil_kg["No %"] = (vil_kg["no"] / vil_kg["Total"] * 100).round(2)
    
    gt_s = vil_kg[["yes", "no", "Total"]].sum()
    gt_row = pd.DataFrame([{
        "village": "Grand Total", "yes": gt_s["yes"], "Yes %": round(gt_s["yes"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "no": gt_s["no"], "No %": round(gt_s["no"]/gt_s["Total"]*100, 2) if gt_s["Total"] > 0 else 0,
        "Total": gt_s["Total"]
    }])
    vil_kg_table = pd.concat([vil_kg, gt_row], ignore_index=True)
    vil_kg_display = vil_kg_table.rename(columns={"yes": Y_L, "no": N_L, "Yes %": f"{Y_L} %", "No %": f"{N_L} %"})
    
    # Legend Professional Fix: Use percentages for Y-axis
    vil_kg["HH with KG %"] = vil_kg["Yes %"]
    vil_kg["HH without KG %"] = (vil_kg["no"] / vil_kg["Total"] * 100).round(2)
    fig_vil = px.bar(vil_kg.sort_values("Yes %", ascending=False), x="village", y=["HH with KG %", "HH without KG %"], title="20.2. Kitchen Gardens - Village level", barmode="group", color_discrete_sequence=COLORS)
    save_and_render("20.2", "Kitchen Gardens - Village level", clean_df(vil_kg_display), fig_vil)

# ==========================================
# MAIN APP FLOW
# ==========================================

def clean_underscore_values(val):
    if pd.isna(val): return val
    s = str(val).replace('_', ' ')
    # Preserve acronyms
    upper_s = s.strip().upper()
    if upper_s in ["SC", "ST", "OBC", "GP", "HH", "LULC", "KG"]:
        return upper_s
    return s.strip().title()

def main():
    st.markdown("""
        <div class='main-header'>
            <h1>🌊 RWI Baseline Report Generator</h1>
            <p>Comprehensive Indicator-wise Analysis Suite</p>
        </div>
    """, unsafe_allow_html=True)
    
    setup_dirs()
    
    with st.sidebar:
        st.header("📂 Data Integration")
        st.info("Upload both sheets to combine social and plot data for the full report.")
        file_hh = st.file_uploader("1. Household Data (Excel)", type=["xlsx"])
        file_plot = st.file_uploader("2. Plot Data (Excel)", type=["xlsx"])
        
    if file_hh and file_plot:
        # Inline Comment: Reading data using the same engine as the original Colab script
        df_hh_all = pd.read_excel(file_hh, engine="openpyxl", sheet_name="Sheet1")
        df_plot_all = pd.read_excel(file_plot, engine="openpyxl", sheet_name="Sheet1")
        
        # Clean underscore in column values for proper text representation
        for df_temp in [df_hh_all, df_plot_all]:
            for col in df_temp.columns:
                if col not in ["income_sources", "type_Irrigation_sources", "Irrigation_sources"]:
                    if df_temp[col].dtype == 'object':
                        df_temp[col] = df_temp[col].apply(clean_underscore_values)
        
        # Inline Comment: Filtering by GP as the original script does for 'Maddivarigondi'
        gps = df_hh_all['gp'].unique()
        selected_gp = st.selectbox("Filter by Gram Panchayat", gps, index=list(gps).index("Maddivarigondi") if "Maddivarigondi" in gps else 0)
        
        df_hh = df_hh_all[df_hh_all['gp'] == selected_gp]
        df_plot = df_plot_all[df_plot_all['gp'] == selected_gp]
        
        if st.button("🚀 Run Total Analysis"):
            setup_dirs()
            st.session_state['run'] = True

    if file_hh and file_plot and st.session_state.get('run', False):
        # Progress Bar
        progress = st.progress(0)
        status = st.empty()
        
        # Pipeline of all ported sections
        pipeline = [
            (run_3_households, "3. HOUSEHOLDS", df_hh),
            (run_4_social_composition, "4. SOCIAL COMPOSITION", df_hh),
            (run_5_occupations, "5. OCCUPATIONS", df_hh),
            (run_6_income_sources, "6. INCOME SOURCES", df_hh),
            (run_7_job_cards, "7. Job Cards", df_hh),
            (run_8_head_of_households, "8. Head of the Household", df_hh),
            (run_9_migration, "9. MIGRATION", df_hh),
            (run_10_institutions, "10. INSTITUTIONS", df_hh),
            (run_11_land_ownership, "11. LAND OWNERSHIP", df_hh),
            (run_12_land_leased, "12. LAND LEASED IN AND LAND LEASED OUT", df_hh),
            (run_13_livestock, "13. LIVESTOCK", df_hh),
            (run_14_soil_types, "14. SOILS", df_plot),
            (run_15_irrigation, "15. IRRIGATION SOURCES", df_plot, df_hh),
            (run_16_lulc, "16. LULC WISE EXTENT", df_plot),
            (run_17_cropping_systems, "17. CROPPING SYSTEMS", df_plot),
            (run_18_cropping_methods, "18. CROPPING METHODS", df_plot),
            (run_19_fallow_lands, "19. FALLOW LANDS", df_plot),
            (run_20_kitchen_gardens, "20. KITCHEN GARDENS", df_hh)
        ]
        
        for i, item in enumerate(pipeline):
            func, name = item[0], item[1]
            args = item[2:]
            status.text(f"Processing Section {name}...")
            func(*args)
            progress.progress((i + 1) / len(pipeline))
        
        status.text("Analysis Complete!")
        
        # Results ZIP
        st.markdown("---")
        zip_path = "RWI_Total_Analysis.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(OUTPUT_DIR):
                for file in files:
                    zipf.write(os.path.join(root, file), arcname=file)
        
        with open(zip_path, "rb") as f:
            st.download_button("📥 Download Total Analysis Package (ZIP)", f, "RWI_Total_Analysis.zip", "application/zip")
        
        # --- Word Report Generation ---
        st.markdown("### 📄 Generate Comprehensive Word Report")
        if st.button("📝 Compile All Sections into Word Doc"):
            doc = Document()
            doc.add_heading(f"RWI Baseline Analysis Report - {selected_gp}", 0)
            
            for section in pipeline:
                func, section_title = section[0], section[1]
                # We need to find all sub-sections (e.g. 3.1, 3.2) for this main section (e.g. 3. HOUSEHOLDS)
                # Since section_id is saved in the filename, we can scan OUTPUT_DIR
                prefix = section_title.split('.')[0]
                
                # Sort files to maintain order (3.1, 3.2, etc)
                files = sorted([f for f in os.listdir(OUTPUT_DIR) if f.startswith(prefix) and f.endswith('.xlsx')])
                
                for excel_file in files:
                    base_id = excel_file.replace('.xlsx', '')
                    png_file = f"{base_id}.png"
                    
                    # Add Section Header in Doc
                    doc.add_heading(f"Section {base_id}", level=1)
                    
                    # Add Image
                    img_path = os.path.join(OUTPUT_DIR, png_file)
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6))
                    
                    # Add Table
                    df_temp = pd.read_excel(os.path.join(OUTPUT_DIR, excel_file))
                    table = doc.add_table(rows=1, cols=len(df_temp.columns))
                    table.style = 'Table Grid'
                    
                    # Header Row
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(df_temp.columns):
                        hdr_cells[i].text = str(col)
                    
                    # Data Rows
                    for _, row in df_temp.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)
                    
                    doc.add_page_break()

            # Save to BytesIO for Streamlit download
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.download_button(
                label="📥 Download Full Report (Word .docx)",
                data=doc_io,
                file_name=f"RWI_Full_Report_{selected_gp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    else:
        st.warning("Please upload both **Household** and **Plot** data sheets in the sidebar to begin.")

if __name__ == "__main__":
    main()
